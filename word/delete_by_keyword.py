#!/usr/bin/python
# -*- coding:UTF-8 -*-
'''
@author:蒋雨航
@file:delete_by_keyword
@time:2023/08/10
@Description:
'''
import os
from docx import Document
import win32com.client as win32


def docx_delete_by_keywords(file_path, key):
    # 打开文档
    try:
        doc = Document(file_path)
    except:
        print(f"docx文档-{file_path}打开失败")
    # 删除内容为"key"的字符
    docx_del(doc, key)

    # 保存修改后的文档
    doc.save(file_path)


def docx_del(doc, key):
    # 段落
    for paragraph in doc.paragraphs:
        for keyword in keywords:
            if keyword in paragraph.text:
                runs = paragraph.runs
                for run in runs:
                    if keyword in run.text:
                        run.text = run.text.replace(keyword, "")
    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for keyword in keywords:
                        if keyword in paragraph.text:
                            runs = paragraph.runs
                            for run in runs:
                                if keyword in run.text:
                                    run.text = run.text.replace(keyword, "")
    # 页眉
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for keyword in keywords:
                if keyword in paragraph.text:
                    runs = paragraph.runs
                    for run in runs:
                        if keyword in run.text:
                            run.text = run.text.replace(keyword, "")

    # 处理页脚
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            for keyword in keywords:
                if keyword in paragraph.text:
                    runs = paragraph.runs
                    for run in runs:
                        if keyword in run.text:
                            run.text = run.text.replace(keyword, "")


def get_directories(directory):
    directories = [name for name in os.listdir(directory) if os.path.isdir(os.path.join(directory, name))]
    return directories


def doc_delete_by_keywords(file_path, characters):
    try:
        # 打开文档
        doc = word_app.Documents.Open(file_path)
    except:
        print(f"doc文档-{file_path}打开失败")

    # 设置查找替换参数
    find_str = characters
    replace_str = ''

    # 执行查找替换操作
    word_app.Selection.Find.ClearFormatting()
    word_app.Selection.Find.Replacement.ClearFormatting()
    word_app.Selection.Find.Execute(find_str, False, False, False, False, False, True, 1, True, replace_str, 2)

    # 处理页眉和页脚
    for section in doc.Sections:
        # 处理页眉
        for header in section.Headers:
            for paragraph in header.Range.Paragraphs:
                for keyword in keywords:
                    if keyword in paragraph.Range.Text:
                        paragraph.Range.Text = paragraph.Range.Text.replace(keyword, "")

        # 处理页脚
        for footer in section.Footers:
            for paragraph in footer.Range.Paragraphs:
                for keyword in keywords:
                    if keyword in paragraph.Range.Text:
                        paragraph.Range.Text = paragraph.Range.Text.replace(keyword, "")

    # 保存并关闭文档
    doc.Save()
    doc.Close()


def convert_to_docx(doc_path):
    # save_path=os.path.join(os.path.dirname(doc_path),"2docx")
    # file_name = os.path.splitext(os.path.basename(doc_path))[0]
    # 打开原始文档
    doc = word_app.Documents.Open(doc_path)
    # 生成新的文件名
    docx_path = os.path.splitext(doc_path)[0] + '.docx'

    # 将原始文档另存为docx格式
    doc.SaveAs2(docx_path, FileFormat=16)
    # 关闭原始文档
    doc.Close()

    return docx_path


if __name__ == '__main__':
    # 创建Word应用程序对象
    word_app = win32.gencache.EnsureDispatch('Word.Application')

    directory = r''
    keywords = ""
    for d in os.listdir(directory):
        file_path = os.path.join(directory, d)
        if file_path.lower().endswith('.docx'):
            # docx文件
            docx_delete_by_keywords(file_path, keywords)
        elif file_path.lower().endswith('.doc'):
            # doc文件
            doc_delete_by_keywords(file_path, keywords)
